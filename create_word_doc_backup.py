from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Apex colors
APEX_NAVY = RGBColor(0, 32, 96)
APEX_BLUE = RGBColor(0, 102, 204)
APEX_BRIGHT_BLUE = RGBColor(59, 130, 246)
APEX_SKY_BLUE = RGBColor(0, 176, 240)
APEX_AMETHYST = RGBColor(128, 44, 192)
APEX_AMETHYST_PINK = RGBColor(236, 0, 117)
APEX_GREEN = RGBColor(40, 167, 69)
APEX_GOLD = RGBColor(251, 191, 36)
APEX_CHARCOAL = RGBColor(45, 55, 72)

def add_shading(cell, color):
    """Add background color to table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_styled_paragraph(doc, text, font_size=11, bold=False, color=APEX_CHARCOAL, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Create a styled paragraph"""
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def create_heading(doc, text, level=1, color=APEX_NAVY):
    """Create a styled heading"""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = color
    heading.runs[0].font.name = 'Calibri'
    return heading

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# PAGE 1: Cover Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('APEX CLEARING')
run.font.size = Pt(48)
run.font.bold = True
run.font.color.rgb = APEX_NAVY

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Model Marketplace')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = APEX_AMETHYST_PINK

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('Business Framework & Implementation Guide')
run.font.size = Pt(18)
run.font.color.rgb = APEX_BLUE

doc.add_paragraph()
doc.add_paragraph()

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('Version 1.1 | March 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# PAGE 2: What is Model Marketplace
create_heading(doc, 'What is Model Marketplace?', 1, APEX_NAVY)

add_heading_with_color(doc, 'What is Model Marketplace?', 2, APEX_BLUE)

para = doc.add_paragraph()
para.add_run("Model Marketplace is Apex's investment model distribution platform integrated within Ascend Rebalancer, allowing advisors to subscribe to professionally managed models from leading investment management providers. Models generate trade proposals based on market activity that realign investment portfolios according to a defined set of holdings and their weight. Subscribed models receive ").font.size = Pt(11)
bold_run = para.add_run("automatic updates and adjustments")
bold_run.font.size = Pt(11)
bold_run.font.bold = True
para.add_run(" from the provider, enabling efficient model updates, allocation changes, and compliance monitoring.").font.size = Pt(11)

add_heading_with_color(doc, 'Key Value Proposition', 3, APEX_NAVY)

add_bullet_point(doc, 'Access to institutional-grade investment models from top providers')
add_bullet_point(doc, 'Automatic updates from providers when models are rebalanced or adjusted')
add_bullet_point(doc, 'Seamless integration with Apex Rebalancer for portfolio construction and ongoing management')
add_bullet_point(doc, 'Simplified subscription and implementation process')
add_bullet_point(doc, 'No subscription fees for clients (free access to marketplace)')
add_bullet_point(doc, 'Available models support efficient portfolio management with automatic weight adjustments')

add_heading_with_color(doc, 'Current Model Providers', 3, APEX_NAVY)

providers_table = doc.add_table(rows=5, cols=2)
providers_table.style = 'Light Grid Accent 1'

providers_data = [
    ('Aptis', 'LIVE'),
    ('Franklin Templeton', 'LIVE'),
    ('Timco and Sachs', 'LIVE'),
    ('State Street', 'COMING SOON')
]

# Header row
header_cells = providers_table.rows[0].cells
header_cells[0].text = 'Provider'
header_cells[1].text = 'Status'
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '002060')
    cell._element.get_or_add_tcPr().append(shading)

# Data rows
for i, (provider, status) in enumerate(providers_data, start=1):
    row = providers_table.rows[i]
    row.cells[0].text = provider
    row.cells[1].text = status

add_heading_with_color(doc, 'Current Market Size', 3, APEX_NAVY)

stats_table = doc.add_table(rows=4, cols=2)
stats_table.style = 'Light Grid Accent 1'

stats_data = [
    ('Active Rebalancer Clients', '38'),
    ('Accounts on Rebalancer', '~50,000'),
    ('Model Providers Live', '4')
]

# Header
stats_header = stats_table.rows[0].cells
stats_header[0].text = 'Metric'
stats_header[1].text = 'Count'
for cell in stats_header:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '0066CC')
    cell._element.get_or_add_tcPr().append(shading)

# Data
for i, (metric, count) in enumerate(stats_data, start=1):
    row = stats_table.rows[i]
    row.cells[0].text = metric
    row.cells[1].text = count

add_page_break(doc)

# ============================================================
# CLIENT ELIGIBILITY
# ============================================================

add_heading_with_color(doc, '1. Client Eligibility & Prerequisites', 1, APEX_NAVY)

add_heading_with_color(doc, 'Who Can Use Model Marketplace?', 2, APEX_BLUE)

# Required section
req_heading = doc.add_paragraph('✓ Required')
req_heading.runs[0].font.size = Pt(14)
req_heading.runs[0].font.bold = True
req_heading.runs[0].font.color.rgb = APEX_BLUE
req_heading.space_before = Pt(12)
req_heading.space_after = Pt(6)

add_bullet_point(doc, 'Must be an Ascend client (not available on Classic platform)', APEX_BLUE)
add_bullet_point(doc, 'Must be an active Rebalancer client', APEX_BLUE)
add_bullet_point(doc, 'Correspondent must be configured for Model Marketplace access', APEX_BLUE)

# Not Eligible section
not_req_heading = doc.add_paragraph('✗ Not Eligible')
not_req_heading.runs[0].font.size = Pt(14)
not_req_heading.runs[0].font.bold = True
not_req_heading.runs[0].font.color.rgb = APEX_CHARCOAL
not_req_heading.space_before = Pt(20)
not_req_heading.space_after = Pt(6)

add_bullet_point(doc, 'Classic platform clients (until Ascend migration complete)', APEX_CHARCOAL)
add_bullet_point(doc, 'Non-Rebalancer clients', APEX_CHARCOAL)
add_bullet_point(doc, 'Clients mid-migration (unless willing to rebalance partial accounts)', APEX_CHARCOAL)

add_page_break(doc)

# ============================================================
# PROCESS FLOW
# ============================================================

add_heading_with_color(doc, '2. Model Marketplace Process Flow', 1, APEX_NAVY)

steps = [
    ('Step 1: Initial Access Setup', 'WHO: Apex Onboarding Team\nWHAT: Enable Model Marketplace in Client Configurator\nWHERE: Experience Settings → Enable Model Marketplace'),
    ('Step 2: User Permissions', 'WHO: Client Administrator\nWHAT: Grant subscribe permissions to authorized users\nACCESS: View vs. Subscribe rights'),
    ('Step 3: Client Discovery', 'WHO: End Client/Advisor\nWHAT: Sign in → Trading → Model Marketplace\nACCEPT: Apex Terms & Conditions (one-time)'),
    ('Step 4: Model Evaluation', 'Review holdings, target weights, summary, resources, fact sheets\nCOLUMNS: Provider | Model | Investment Vehicle | Status'),
    ('Step 5: Model Subscription', 'Click Subscribe → Accept Provider T&Cs\nRESULT: Status changes to SUBSCRIBED\nButton: Green → Red (Unsubscribe)'),
    ('Step 6: Rebalancer Configuration', 'CRITICAL: Verify asset class assignments\nCreate goals → Assign to accounts'),
    ('Step 7: Portfolio Rebalancing', 'Initiate rebalance → Review proposals\nReview compliance disclosures\nACTION: Submit trades for execution'),
    ('Step 8: Unsubscribe (Optional)', 'Stops automatic updates\nRetains last subscribed weights\nTO DISCONTINUE: Delete model from Rebalancer')
]

for title, content in steps:
    step_heading = doc.add_paragraph(title)
    step_heading.runs[0].font.size = Pt(13)
    step_heading.runs[0].font.bold = True
    step_heading.runs[0].font.color.rgb = APEX_NAVY
    step_heading.space_before = Pt(10)
    step_heading.space_after = Pt(6)

    step_content = doc.add_paragraph(content)
    step_content.runs[0].font.size = Pt(11)
    step_content.runs[0].font.color.rgb = APEX_CHARCOAL
    step_content.space_after = Pt(8)

# Add warning box
warning = doc.add_paragraph()
warning_run = warning.add_run('⚠️ KEY LIMITATION: ')
warning_run.font.bold = True
warning_run.font.color.rgb = RGBColor(239, 68, 68)  # Red
warning_run.font.size = Pt(11)
warning.add_run('Apex does NOT automatically execute trades. Automatic scheduled rebalancing planned for Q3 2026. Clients must currently review and submit all trade orders manually.').font.size = Pt(11)

add_page_break(doc)

# Save document
doc.save(r'C:\Users\uariyasena\Projects\Model Market Place\Model_Marketplace_Framework.docx')
print("Word document created successfully!")
