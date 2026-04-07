from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Apex colors
APEX_NAVY = RGBColor(0, 32, 96)
APEX_BLUE = RGBColor(0, 102, 204)
APEX_BRIGHT_BLUE = RGBColor(59, 130, 246)
APEX_SKY_BLUE = RGBColor(0, 176, 240)
APEX_AMETHYST = RGBColor(128, 44, 192)
APEX_AMETHYST_PINK = RGBColor(236, 0, 117)
APEX_GREEN = RGBColor(40, 167, 69)
APEX_GOLD = RGBColor(251, 191, 36)
APEX_CHARCOAL = RGBColor(45, 55, 72)

def add_shading(cell, color):
    """Add background color to table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create document
doc = Document()

# PAGE 1: Cover Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('APEX CLEARING')
run.font.size = Pt(48)
run.font.bold = True
run.font.color.rgb = APEX_NAVY

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Model Marketplace')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = APEX_AMETHYST_PINK

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('Business Framework & Implementation Guide')
run.font.size = Pt(18)
run.font.color.rgb = APEX_BLUE

for _ in range(3):
    doc.add_paragraph()

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('Version 1.1 | March 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# PAGE 2: What is Model Marketplace
heading = doc.add_heading('What is Model Marketplace?', 1)
heading.runs[0].font.color.rgb = APEX_NAVY

p = doc.add_paragraph('Model Marketplace is Apex Clearing\'s centralized platform that enables investment advisors to discover, subscribe to, and implement third-party investment models directly within Ascend Workstation.')

heading2 = doc.add_heading('Core Functionality', 2)
heading2.runs[0].font.color.rgb = APEX_AMETHYST

doc.add_paragraph('Provides a curated marketplace of vetted investment models from institutional providers', style='List Bullet')
doc.add_paragraph('Enables one-click subscription to professional asset allocation strategies', style='List Bullet')
doc.add_paragraph('Automatically delivers model updates from providers to subscribed clients', style='List Bullet')
doc.add_paragraph('Integrates subscribed models directly into Apex Rebalancer as ready-to-use portfolio goals', style='List Bullet')

heading2 = doc.add_heading('How Model Marketplace Connects to Rebalancer', 2)
heading2.runs[0].font.color.rgb = APEX_AMETHYST

p = doc.add_paragraph('Model Marketplace and Rebalancer work as an integrated ecosystem. When an advisor subscribes to a model in Model Marketplace, that model automatically becomes available as a \'Goal\' in Rebalancer.')

# Workflow table
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'Step'
header_cells[1].text = 'Action'
header_cells[2].text = 'Result'
for cell in header_cells:
    add_shading(cell, '0066CC')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True

workflow_data = [
    ['1. Subscribe', 'Subscribe in Model Marketplace', 'Model appears in Rebalancer Goals list'],
    ['2. Configure', 'Assign goal to client accounts', 'Portfolio linked to model allocation'],
    ['3. Rebalance', 'Generate buy/sell orders', 'Portfolio matches model targets'],
    ['4. Auto-Update', 'Provider updates model', 'Rebalancer goal updates automatically']
]

for i, (step, action, result) in enumerate(workflow_data, 1):
    cells = table.rows[i].cells
    cells[0].text = step
    cells[1].text = action
    cells[2].text = result

heading2 = doc.add_heading('Current Status (Q1 2026)', 2)
heading2.runs[0].font.color.rgb = APEX_AMETHYST

p = doc.add_paragraph()
run = p.add_run('Total Providers: ')
run.font.bold = True
p.add_run('4 (Aptis, Franklin Templeton, Timco and Sachs, State Street coming soon)')

p = doc.add_paragraph()
run = p.add_run('Models Available: ')
run.font.bold = True
p.add_run('12+')

p = doc.add_paragraph()
run = p.add_run('Target Clients: ')
run.font.bold = True
p.add_run('38 existing Rebalancer firms')

p = doc.add_paragraph()
run = p.add_run('Phase: ')
run.font.bold = True
p.add_run('Marketing launch & press release (Q2 2026)')

doc.add_page_break()

# PAGE 3: 8-Step Process Part 1
heading = doc.add_heading('How Users Access & Use Model Marketplace', 1)
heading.runs[0].font.color.rgb = APEX_NAVY

steps_part1 = [
    ('Step 1: Initial Access Setup', 'WHO: Apex Onboarding Team / Client Configurator Team', 'WHAT: Enable Model Marketplace at correspondent level', 'WHERE: Client Configurator → Experience Settings', 'TIMELINE: During client onboarding'),
    ('Step 2: User Permission Configuration', 'WHO: Client Administrator / Apex Client Services', 'WHAT: Grant subscribe permissions to specific users', 'WHERE: User management interface', 'TIMELINE: Before first use'),
    ('Step 3: Client Acceptance & Discovery', 'WHO: End Client/Advisor', 'WHAT: Sign in to Ascend → Trading → Model Marketplace', 'WHERE: Accept Apex Terms & Conditions (one-time)', 'TIMELINE: 5-10 minutes'),
    ('Step 4: Model Evaluation', 'WHO: Client/Advisor', 'WHAT: Review model details: Holdings, Target Weights, Summary', 'WHERE: Filter by Provider, Model, Investment Vehicle, Status', 'TIMELINE: Variable - depends on research depth')
]

for title, who, what, where, timeline in steps_part1:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = APEX_AMETHYST_PINK

    doc.add_paragraph(who, style='List Bullet')
    doc.add_paragraph(what, style='List Bullet')
    doc.add_paragraph(where, style='List Bullet')
    doc.add_paragraph(timeline, style='List Bullet')

doc.add_page_break()

# PAGE 4: 8-Step Process Part 2
heading = doc.add_heading('How Users Access & Use Model Marketplace (continued)', 1)
heading.runs[0].font.color.rgb = APEX_NAVY

steps_part2 = [
    ('Step 5: Model Subscription', 'WHO: Authorized Client User (with subscribe permission)', 'WHAT: Click Subscribe → Accept Provider-Specific T&Cs', 'WHERE: Status changes: AVAILABLE → SUBSCRIBED', 'TIMELINE: 2-3 minutes per model'),
    ('Step 6: Post-Subscription Configuration in Rebalancer', 'WHO: Client/Advisor', 'WHAT: CRITICAL - Verify asset class assignments for each security', 'WHERE: Rebalancer → Settings → Models → Create Goals', 'TIMELINE: 15-30 minutes for first model'),
    ('Step 7: Portfolio Rebalancing', 'WHO: Client/Advisor', 'WHAT: Initiate rebalance → Review buy/sell orders → Submit trades', 'WHERE: Review compliance disclosures before submission', 'TIMELINE: Variable based on portfolio size'),
    ('Step 8: Unsubscribe (Optional)', 'WHO: Client/Advisor', 'WHAT: Stops automatic updates but retains last subscribed weights', 'WHERE: To fully discontinue: Delete model from Rebalancer', 'TIMELINE: Immediate')
]

for title, who, what, where, timeline in steps_part2:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = APEX_AMETHYST_PINK

    doc.add_paragraph(who, style='List Bullet')
    doc.add_paragraph(what, style='List Bullet')
    doc.add_paragraph(where, style='List Bullet')
    doc.add_paragraph(timeline, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('⚠️ KEY LIMITATION: ')
run.font.bold = True
run.font.color.rgb = RGBColor(255, 0, 0)
run = p.add_run('Apex does NOT automatically execute trades. Automatic scheduled rebalancing planned for Q3 2026.')
run.font.color.rgb = RGBColor(255, 0, 0)

doc.add_page_break()

# PAGE 5: Success Metrics
heading = doc.add_heading('Success Metrics & Targets (2026)', 1)
heading.runs[0].font.color.rgb = APEX_NAVY

heading2 = doc.add_heading('Adoption Targets', 2)
heading2.runs[0].font.color.rgb = APEX_AMETHYST

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'Metric'
header_cells[1].text = 'Current (Q1)'
header_cells[2].text = 'Target (2026)'
for cell in header_cells:
    add_shading(cell, '002060')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True

adoption_data = [
    ['Total Rebalancer Clients', '500', '38 enabled'],
    ['MM Enabled Clients', '38', '30 (80%)'],
    ['Users with Subscribe Permission', 'TBD', '100+'],
    ['Active Model Subscriptions', 'Growing', '200+']
]

for i, (metric, current, target) in enumerate(adoption_data, 1):
    cells = table.rows[i].cells
    cells[0].text = metric
    cells[1].text = current
    cells[2].text = target

heading2 = doc.add_heading('Business Impact Targets', 2)
heading2.runs[0].font.color.rgb = APEX_AMETHYST

table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'Metric'
header_cells[1].text = 'Target'
for cell in header_cells:
    add_shading(cell, '802CC0')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True

impact_data = [
    ['AUM in Models', '$500M+ by EOY 2026'],
    ['Revenue per Model', '$50K average annual'],
    ['Client Retention Impact', '+15% for MM users'],
    ['New Correspondent Wins', 'MM cited in 25% of pitches']
]

for i, (metric, target) in enumerate(impact_data, 1):
    cells = table.rows[i].cells
    cells[0].text = metric
    cells[1].text = target

doc.add_page_break()

# PAGE 6: Provider Integration & Client Communication
heading = doc.add_heading('Provider Integration & Client Communication', 1)
heading.runs[0].font.color.rgb = APEX_NAVY

heading2 = doc.add_heading('Current Providers', 2)
heading2.runs[0].font.color.rgb = APEX_AMETHYST

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'Provider'
header_cells[1].text = 'Models'
header_cells[2].text = 'Integration'
for cell in header_cells:
    add_shading(cell, '802CC0')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True

provider_data = [
    ['Aptis', 'ETF models', 'LIVE'],
    ['Franklin Templeton', 'ETF models', 'LIVE - SFTP (daily feeds)'],
    ['Timco and Sachs', 'Investment models', 'LIVE'],
    ['State Street', 'TBD', 'COMING SOON']
]

for i, (provider, models, integration) in enumerate(provider_data, 1):
    cells = table.rows[i].cells
    cells[0].text = provider
    cells[1].text = models
    cells[2].text = integration

heading2 = doc.add_heading('Communication Strategy', 2)
heading2.runs[0].font.color.rgb = APEX_AMETHYST

doc.add_paragraph('Direct Outreach: Personalized emails to 38 existing Rebalancer clients', style='List Bullet')
doc.add_paragraph('Documentation: Service Center articles, video tutorials', style='List Bullet')
doc.add_paragraph('Press Release: Q2 2026 announcement targeting industry media', style='List Bullet')
doc.add_paragraph('Internal Enablement: Sales team training, demo environment', style='List Bullet')

doc.add_page_break()

# PAGE 7: Rebalancer Integration
heading = doc.add_heading('Rebalancer Capabilities & Integration', 1)
heading.runs[0].font.color.rgb = APEX_NAVY

heading2 = doc.add_heading('What is Rebalancer?', 2)
heading2.runs[0].font.color.rgb = APEX_AMETHYST

p = doc.add_paragraph('Rebalancer is Apex\'s portfolio management tool that enables advisors to create target asset allocations (goals), assign them to accounts, and generate buy/sell orders. Model Marketplace subscriptions automatically populate as available goals within Rebalancer.')

heading2 = doc.add_heading('Core Features', 2)
heading2.runs[0].font.color.rgb = APEX_AMETHYST

table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'Feature'
header_cells[1].text = 'Description'
for cell in header_cells:
    add_shading(cell, '002060')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True

features = [
    ['Goals Management', 'Create custom allocations or use subscribed models'],
    ['Account Assignment', 'Link goals to individual or household accounts'],
    ['Drift Monitoring', 'Track portfolio deviation from target allocation'],
    ['Order Generation', 'Automated buy/sell recommendations'],
    ['Tax Optimization', 'Tax-loss harvesting and gain/loss awareness'],
    ['Compliance Checks', 'Pre-trade compliance rules and disclosures']
]

for i, (feature, description) in enumerate(features, 1):
    cells = table.rows[i].cells
    cells[0].text = feature
    cells[1].text = description

p = doc.add_paragraph()
run = p.add_run('⚠️ Critical: ')
run.font.bold = True
run.font.color.rgb = RGBColor(255, 0, 0)
run = p.add_run('After subscribing, clients MUST verify asset class assignments for each security in the model.')
run.font.color.rgb = RGBColor(255, 0, 0)

doc.add_page_break()

# PAGE 8: Roadmap & Contacts
heading = doc.add_heading('Roadmap & Key Contacts', 1)
heading.runs[0].font.color.rgb = APEX_NAVY

heading2 = doc.add_heading('2026 Roadmap', 2)
heading2.runs[0].font.color.rgb = APEX_AMETHYST

table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'Quarter'
header_cells[1].text = 'Initiative'
header_cells[2].text = 'Status'
for cell in header_cells:
    add_shading(cell, '0066CC')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True

roadmap = [
    ['Q2 2026', 'Marketing Launch & Press Release', 'IN PROGRESS'],
    ['Q2 2026', 'Enhanced Provider Onboarding', 'IN PROGRESS'],
    ['Q3 2026', 'Automatic Scheduled Rebalancing', 'PLANNED'],
    ['Q3 2026', 'BD Resource for Provider Acquisition', 'PLANNED'],
    ['Q4 2026', 'Advanced Filtering (ESG, risk)', 'PLANNED'],
    ['Q4 2026', 'Performance Analytics Dashboard', 'PLANNED']
]

for i, (quarter, initiative, status) in enumerate(roadmap, 1):
    cells = table.rows[i].cells
    cells[0].text = quarter
    cells[1].text = initiative
    cells[2].text = status

p = doc.add_paragraph()
run = p.add_run('Top Priority: ')
run.font.bold = True
run.font.color.rgb = APEX_BLUE
p.add_run('Q3 2026 will introduce automatic scheduled rebalancing where clients can set recurring schedules (weekly, monthly, quarterly).')

heading2 = doc.add_heading('Key Contacts', 2)
heading2.runs[0].font.color.rgb = APEX_AMETHYST

table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'Role'
header_cells[1].text = 'Contact'
for cell in header_cells:
    add_shading(cell, '002060')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True

contacts = [
    ['Product Owner', 'Neema Raphael - neema.raphael@apexclearing.com'],
    ['Technical Support', 'Ascend Support - support@apexclearing.com'],
    ['Sales Inquiries', 'Business Development - bd@apexclearing.com'],
    ['Provider Partnerships', 'BD Team - partnerships@apexclearing.com']
]

for i, (role, contact) in enumerate(contacts, 1):
    cells = table.rows[i].cells
    cells[0].text = role
    cells[1].text = contact

doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Document Version: 1.1 | Last Updated: March 2026\nConfidential: For internal Apex use and authorized correspondents only')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

# Save document
doc.save('Model_Marketplace_Framework.docx')
print("8-page Word document created successfully: Model_Marketplace_Framework.docx")
print("Location: C:\\Users\\uariyasena\\Projects\\Model Market Place\\Model_Marketplace_Framework.docx")
